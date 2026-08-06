# app.py
# Wanderly - Local Travel Guide Chatbot (TRV-03)
# Core Concept: Search-enabled Agent
# Key Features: Nearby attractions, restaurants, hotels, budget itineraries,
#               best time to visit, downloadable itineraries, place photos,
#               Hinglish-friendly replies.
# Stack: LangChain (create_agent) + Groq (LLM) + OpenStreetMap (places search)
#        + Open-Meteo (free climate data) + Wikipedia (free place photos)
# All of these are free - no credit card needed anywhere. Only Groq needs a
# free API key (entered by the user in the sidebar); everything else is keyless.

import io
from datetime import datetime
from pathlib import Path
from urllib.parse import quote
import base64

import requests
import streamlit as st
from docx import Document
from fpdf import FPDF
from langchain.agents import create_agent
from langchain_groq import ChatGroq

HEADERS = {"User-Agent": "wanderly-travel-chatbot/1.0"}
MODEL_NAME = "openai/gpt-oss-120b"


# =========================================================
# SHARED HELPERS (not tools themselves)
# =========================================================

def geocode(location):
    """Turn a place name into (lat, lon, resolved_name) using free Nominatim. Returns None if not found."""
    try:
        geo = requests.get(
            "https://nominatim.openstreetmap.org/search",
            params={"q": location, "format": "json", "limit": 1},
            headers=HEADERS, timeout=15,
        ).json()
        if not geo:
            return None
        return float(geo[0]["lat"]), float(geo[0]["lon"]), geo[0].get("display_name", location)
    except Exception:
        return None


def overpass_search(lat, lon, tag, radius=3000, limit=5):
    """Query the free Overpass API for named elements matching `tag` near (lat, lon)."""
    query = f'[out:json][timeout:25];(node{tag}(around:{radius},{lat},{lon});); out {limit * 2};'
    try:
        elements = requests.post(
            "https://overpass-api.de/api/interpreter",
            data={"data": query}, headers=HEADERS, timeout=15,
        ).json().get("elements", [])
    except Exception:
        return []
    named = [e for e in elements if e.get("tags", {}).get("name")]
    return named[:limit]


def get_place_image(name):
    """Free, keyless: look up a thumbnail photo for a named place via Wikipedia's summary API."""
    try:
        resp = requests.get(
            f"https://en.wikipedia.org/api/rest_v1/page/summary/{quote(name)}",
            headers=HEADERS, timeout=10,
        )
        if resp.status_code != 200:
            return None
        data = resp.json()
        thumb = data.get("thumbnail", {}).get("source") or data.get("originalimage", {}).get("source")
        return thumb
    except Exception:
        return None


def add_to_gallery(names):
    """Fetch and stash a few place photos for the current turn's chat display."""
    gallery = st.session_state.setdefault("gallery", [])
    for name in names:
        img = get_place_image(name)
        if img:
            gallery.append((name, img))


def format_lines(elements, kind_keys):
    lines = []
    for e in elements:
        name = e["tags"]["name"]
        kind = next((e["tags"][k] for k in kind_keys if e["tags"].get(k)), "place")
        lines.append(f"- {name} ({kind})")
    return lines


# =========================================================
# TOOLS
# Plain python functions with a docstring -> create_agent
# auto-wraps these as tools, same pattern used in the
# reference files (temp_tool, search_latest_info, etc.)
# =========================================================

def find_nearby_attractions(location: str) -> str:
    """Find tourist attractions and points of interest near a given location, for example 'Jaipur, India'."""
    geo = geocode(location)
    if not geo:
        return f"Could not find a location matching '{location}'."
    lat, lon, resolved = geo
    elements = overpass_search(lat, lon, '["tourism"]')
    if not elements:
        return f"No attractions found near {resolved}."
    add_to_gallery([e["tags"]["name"] for e in elements[:4]])
    return f"Attractions near {resolved}:\n" + "\n".join(format_lines(elements, ["tourism"]))


def find_nearby_restaurants(location: str) -> str:
    """Find restaurants and cafes near a given location, for example 'Jaipur, India'."""
    geo = geocode(location)
    if not geo:
        return f"Could not find a location matching '{location}'."
    lat, lon, resolved = geo
    elements = overpass_search(lat, lon, '["amenity"~"restaurant|cafe"]')
    if not elements:
        return f"No restaurants found near {resolved}."
    return f"Restaurants near {resolved}:\n" + "\n".join(format_lines(elements, ["amenity", "cuisine"]))


def find_nearby_hotels(location: str) -> str:
    """Find hotels, guest houses, and hostels near a given location, for example 'Jaipur, India'."""
    geo = geocode(location)
    if not geo:
        return f"Could not find a location matching '{location}'."
    lat, lon, resolved = geo
    elements = overpass_search(lat, lon, '["tourism"~"hotel|guest_house|hostel|motel"]')
    if not elements:
        return f"No hotels found near {resolved}."
    add_to_gallery([e["tags"]["name"] for e in elements[:3]])
    return f"Hotels near {resolved}:\n" + "\n".join(format_lines(elements, ["tourism"]))


def plan_daily_budget(total_budget: float, num_days: int, num_travelers: int = 1) -> str:
    """Split a total trip budget into a per-day, per-person spending plan with a
    suggested food/activities/transport/buffer breakdown."""
    if num_days <= 0 or num_travelers <= 0:
        return "num_days and num_travelers must be greater than zero."
    per_day = total_budget / num_days
    per_day_person = per_day / num_travelers
    return (
        f"Total budget: {total_budget} for {num_days} day(s), {num_travelers} traveler(s).\n"
        f"Per day (all travelers): {per_day:.2f}\n"
        f"Per day, per person: {per_day_person:.2f}\n"
        f"Suggested split per person/day -> "
        f"Food: {per_day_person * 0.4:.2f}, "
        f"Activities: {per_day_person * 0.3:.2f}, "
        f"Transport: {per_day_person * 0.15:.2f}, "
        f"Buffer: {per_day_person * 0.15:.2f}"
    )


def best_time_to_visit(location: str) -> str:
    """Find the best months to visit a location, based on free historical temperature
    and rainfall data, for example 'Jaipur, India'."""
    geo = geocode(location)
    if not geo:
        return f"Could not find a location matching '{location}'."
    lat, lon, resolved = geo
    try:
        resp = requests.get(
            "https://archive-api.open-meteo.com/v1/archive",
            params={
                "latitude": lat, "longitude": lon,
                "start_date": "2024-01-01", "end_date": "2024-12-31",
                "daily": "temperature_2m_mean,precipitation_sum",
                "timezone": "auto",
            },
            headers=HEADERS, timeout=15,
        ).json()
        daily = resp.get("daily", {})
        dates, temps, rain = daily.get("time", []), daily.get("temperature_2m_mean", []), daily.get("precipitation_sum", [])
        if not dates:
            return f"Could not fetch climate data for {resolved}."

        months = {}
        for d, t, r in zip(dates, temps, rain):
            m = int(d.split("-")[1])
            months.setdefault(m, {"temp": [], "rain": []})
            if t is not None:
                months[m]["temp"].append(t)
            if r is not None:
                months[m]["rain"].append(r)

        month_names = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
        lines, scored = [], []
        for m in range(1, 13):
            d = months.get(m, {"temp": [], "rain": []})
            if not d["temp"]:
                continue
            avg_temp = sum(d["temp"]) / len(d["temp"])
            total_rain = sum(d["rain"]) if d["rain"] else 0
            score = -abs(avg_temp - 22) - total_rain * 0.1  # comfortable ~22C, less rain is better
            scored.append((score, month_names[m - 1]))
            lines.append(f"{month_names[m - 1]}: avg {avg_temp:.1f}\u00b0C, {total_rain:.0f}mm rain")

        scored.sort(reverse=True)
        best = ", ".join(name for _, name in scored[:3])
        worst = ", ".join(name for _, name in scored[-2:])
        return (
            f"Climate summary for {resolved} (based on last year):\n" + "\n".join(lines)
            + f"\n\nBest months to visit: {best} (milder temperatures, less rain)"
            + f"\nMonths to avoid: {worst} (extreme heat/cold or heavy rain)"
        )
    except Exception as e:
        return f"Could not fetch climate data: {e}"


def build_itinerary(location: str, num_days: int, total_budget: float = None, num_travelers: int = 1) -> str:
    """Build a day-wise travel itinerary for a location, combining nearby attractions,
    a restaurant suggestion per day, and a hotel suggestion. Also prepares the
    itinerary for download as a PDF or DOCX file. total_budget is optional."""
    geo = geocode(location)
    if not geo:
        return f"Could not find a location matching '{location}'."
    lat, lon, resolved = geo

    attractions = overpass_search(lat, lon, '["tourism"]', limit=max(num_days * 2, 6))
    restaurants = overpass_search(lat, lon, '["amenity"~"restaurant|cafe"]', limit=max(num_days, 3))
    hotels = overpass_search(lat, lon, '["tourism"~"hotel|guest_house|hostel"]', limit=1)

    if not attractions:
        return f"Not enough place data found near {resolved} to build an itinerary."

    days = []
    for d in range(num_days):
        day_attractions = attractions[d * 2:(d * 2) + 2] or attractions[:2]
        day_restaurant = restaurants[d % len(restaurants)] if restaurants else None
        days.append({
            "day": d + 1,
            "attractions": [a["tags"]["name"] for a in day_attractions],
            "restaurant": day_restaurant["tags"]["name"] if day_restaurant else None,
        })
        add_to_gallery([a["tags"]["name"] for a in day_attractions])

    budget_text = plan_daily_budget(total_budget, num_days, num_travelers) if total_budget else None

    itinerary = {
        "location": resolved,
        "num_days": num_days,
        "hotel": hotels[0]["tags"]["name"] if hotels else None,
        "days": days,
        "budget_text": budget_text,
        "generated": datetime.now().strftime("%d %b %Y"),
    }
    st.session_state["last_itinerary"] = itinerary

    lines = [f"Itinerary for {resolved} ({num_days} day(s)):"]
    if itinerary["hotel"]:
        lines.append(f"Suggested stay: {itinerary['hotel']}")
    for day in days:
        lines.append(f"\nDay {day['day']}:")
        lines += [f"  - Visit {a}" for a in day["attractions"]]
        if day["restaurant"]:
            lines.append(f"  - Eat at {day['restaurant']}")
    if budget_text:
        lines.append(f"\n{budget_text}")
    lines.append("\n(Download this itinerary as a PDF or DOCX using the buttons below the chat.)")
    return "\n".join(lines)


# =========================================================
# DOCUMENT GENERATION (for the download buttons)
# =========================================================

def itinerary_to_docx(itinerary: dict) -> bytes:
    doc = Document()
    doc.add_heading(f"Travel Itinerary \u2013 {itinerary['location']}", level=1)
    doc.add_paragraph(f"Generated on {itinerary['generated']} | {itinerary['num_days']} day(s)")
    if itinerary.get("hotel"):
        doc.add_paragraph(f"Suggested stay: {itinerary['hotel']}")
    for day in itinerary["days"]:
        doc.add_heading(f"Day {day['day']}", level=2)
        for a in day["attractions"]:
            doc.add_paragraph(f"Visit: {a}", style="List Bullet")
        if day["restaurant"]:
            doc.add_paragraph(f"Eat at: {day['restaurant']}", style="List Bullet")
    if itinerary.get("budget_text"):
        doc.add_heading("Budget Plan", level=2)
        doc.add_paragraph(itinerary["budget_text"])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def itinerary_to_pdf(itinerary: dict) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, f"Travel Itinerary - {itinerary['location']}")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 8, f"Generated on {itinerary['generated']} | {itinerary['num_days']} day(s)")
    if itinerary.get("hotel"):
        pdf.multi_cell(0, 8, f"Suggested stay: {itinerary['hotel']}")
    for day in itinerary["days"]:
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 13)
        pdf.multi_cell(0, 8, f"Day {day['day']}")
        pdf.set_font("Helvetica", "", 11)
        for a in day["attractions"]:
            pdf.multi_cell(0, 7, f"- Visit {a}")
        if day["restaurant"]:
            pdf.multi_cell(0, 7, f"- Eat at {day['restaurant']}")
    if itinerary.get("budget_text"):
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 13)
        pdf.multi_cell(0, 8, "Budget Plan")
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 7, itinerary["budget_text"])
    return bytes(pdf.output())


# =========================================================
# STREAMLIT FRONTEND
# =========================================================

st.set_page_config(page_title="Wanderly - Local Travel Guide Chatbot", page_icon="🧭")


def set_background(image_path: str) -> None:
    """Base64-embed a background image via CSS - works both locally and once deployed."""
    path = Path(image_path)
    if not path.exists():
        return
    encoded = base64.b64encode(path.read_bytes()).decode()
    ext = path.suffix.lstrip(".")
    st.markdown(
        f"""
        <style>
        .stApp {{
            background-image:
                linear-gradient(rgba(10, 20, 20, 0.72), rgba(10, 20, 20, 0.72)),
                url("data:image/{ext};base64,{encoded}");
            background-size: cover;
            background-position: center;
            background-attachment: fixed;
            background-repeat: no-repeat;
        }}
        [data-testid="stSidebar"] {{ background-color: rgba(10, 20, 20, 0.85); }}
        [data-testid="stChatMessage"] {{ background-color: rgba(20, 30, 30, 0.75); border-radius: 12px; }}
        </style>
        """,
        unsafe_allow_html=True,
    )


set_background(str(Path(__file__).parent / "assets" / "background.jpg"))

st.title("🧭 Wanderly — Local Travel Guide Chatbot")
st.caption("Attractions, restaurants, hotels, budgets & best time to visit · LangChain + Groq + OpenStreetMap")

with st.sidebar:
    st.header("🔑 Groq API Key")
    st.caption("Free, no credit card. Get one at console.groq.com/keys")
    groq_api_key = st.text_input("Groq API Key", type="password", placeholder="gsk_...")
    st.markdown("---")
    st.markdown(
        "**Try asking:**\n"
        "- Find attractions near Jaipur, India\n"
        "- Suggest hotels near Amer Fort\n"
        "- Best time to visit Goa\n"
        "- Plan a 3-day itinerary for Jaipur with a $300 budget for 2 people\n"
        "- Jaipur mein ghumne ke liye kya kya hai?"
    )
    if st.button("🔄 Reset conversation"):
        st.session_state.clear()
        st.rerun()

    # --- About section, kept last so it stays at the bottom of the sidebar ---
    st.markdown("---")
    st.caption(
        "ℹ️ **About Wanderly**  \n"
        "A search-enabled travel agent built with LangChain + Groq. It looks up real "
        "attractions, restaurants, and hotels on OpenStreetMap, checks the best months "
        "to visit using free climate data, builds downloadable day-wise itineraries, "
        "and replies in whatever language style you use - English, Hindi, or Hinglish."
    )

if not groq_api_key:
    st.info("👈 Enter your free Groq API key in the sidebar to start chatting.")
    st.stop()

SYSTEM_PROMPT = (
    "You are Wanderly, a friendly local travel guide chatbot. Use your tools to find "
    "real attractions, restaurants, and hotels, to check the best months to visit, to "
    "plan budgets, and to build day-wise itineraries - never invent place names, prices, "
    "or dates; only report what a tool actually returns. Reply in the same language and "
    "style the user writes in: if they write in Hinglish (mixed Hindi-English in Roman "
    "script), reply in Hinglish too; if they write in plain English, reply in English; "
    "if they write in Hindi, reply in Hindi. Keep answers concise and use bullet points "
    "for lists of places."
)

if "agent" not in st.session_state or st.session_state.get("key") != groq_api_key:
    groq_llm = ChatGroq(model=MODEL_NAME, api_key=groq_api_key)
    st.session_state.agent = create_agent(
        model=groq_llm,
        tools=[
            find_nearby_attractions,
            find_nearby_restaurants,
            find_nearby_hotels,
            plan_daily_budget,
            best_time_to_visit,
            build_itinerary,
        ],
        system_prompt=SYSTEM_PROMPT,
    )
    st.session_state.key = groq_api_key

if "messages" not in st.session_state:
    st.session_state.messages = []       # for rendering chat bubbles
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []   # for the agent's "messages" input

for role, text in st.session_state.messages:
    with st.chat_message(role):
        st.markdown(text)

user_input = st.chat_input("Ask about attractions, restaurants, hotels, budget, or best time to visit...")

if user_input:
    st.session_state.messages.append(("user", user_input))
    with st.chat_message("user"):
        st.markdown(user_input)

    pending_history = st.session_state.chat_history + [{"role": "user", "content": user_input}]
    st.session_state["gallery"] = []  # only show photos gathered during THIS turn

    with st.chat_message("assistant"):
        with st.spinner("Thinking..."):
            try:
                response = st.session_state.agent.invoke({"messages": pending_history})
                content = response["messages"][-1].content
                if isinstance(content, list):
                    reply = content[-1].get("text", str(content[-1])) if content else ""
                else:
                    reply = content
                if not reply:
                    reply = "I didn't get a usable response - try rephrasing your question."
                st.session_state.chat_history = pending_history
            except Exception as e:
                reply = (
                    f"Something went wrong calling the model: {e}\n\n"
                    "This is usually a temporary model/tool-calling hiccup - try "
                    "rephrasing your question, or click 'Reset conversation' in the sidebar."
                )
        st.markdown(reply)

        gallery = st.session_state.get("gallery", [])
        if gallery:
            cols = st.columns(min(len(gallery), 4))
            for i, (name, url) in enumerate(gallery[:4]):
                with cols[i % len(cols)]:
                    st.image(url, caption=name, use_container_width=True)

    st.session_state.messages.append(("assistant", reply))

# --- Download section: appears once an itinerary has been built ---
if st.session_state.get("last_itinerary"):
    st.markdown("---")
    st.subheader("📥 Download your itinerary")
    itinerary = st.session_state["last_itinerary"]
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "Download as DOCX", data=itinerary_to_docx(itinerary),
            file_name="itinerary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    with col2:
        st.download_button(
            "Download as PDF", data=itinerary_to_pdf(itinerary),
            file_name="itinerary.pdf", mime="application/pdf",
        )
